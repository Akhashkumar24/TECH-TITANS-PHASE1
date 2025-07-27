# utils/file_handler.py
import os
import shutil
from pathlib import Path
from typing import Optional
from utils.logger import get_logger
import re

# Try to import text extraction libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not installed. PDF text extraction will be limited.")
    print("Install with: pip install PyPDF2")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX text extraction will be limited.")
    print("Install with: pip install python-docx")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("Info: pdfplumber not installed. Using PyPDF2 for PDF extraction.")
    print("For better PDF extraction, install with: pip install pdfplumber")

logger = get_logger(__name__)

class FileHandler:
    def __init__(self):
        self.upload_dir = "uploads/resumes"
        self.ensure_upload_directory()
    
    def ensure_upload_directory(self):
        """Create upload directory if it doesn't exist"""
        Path(self.upload_dir).mkdir(parents=True, exist_ok=True)
    
    def save_resume(self, file_path: str, user_id: int, job_id: int) -> Optional[str]:
        """Save resume file and return the saved path"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"File does not exist: {file_path}")
                return None
            
            # Get file extension
            file_ext = Path(file_path).suffix
            
            # Create new filename
            new_filename = f"resume_{user_id}_{job_id}{file_ext}"
            new_path = os.path.join(self.upload_dir, new_filename)
            
            # Copy file to upload directory
            shutil.copy2(file_path, new_path)
            
            logger.info(f"Resume saved: {new_path}")
            return new_path
            
        except Exception as e:
            logger.error(f"Error saving resume: {e}")
            return None
    
    def validate_file(self, file_path: str) -> bool:
        """Validate file type and size"""
        try:
            if not os.path.exists(file_path):
                return False
            
            # Check file size (10MB limit, increased for better support)
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > 10:
                logger.error(f"File too large: {file_size_mb}MB")
                return False
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            supported_types = ['.pdf', '.docx', '.txt', '.doc']
            if file_ext not in supported_types:
                logger.error(f"Unsupported file type: {file_ext}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating file: {e}")
            return False
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from various file formats"""
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return ""
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                return self._extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_ext}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    logger.info(f"Successfully extracted text from PDF using pdfplumber: {len(text)} characters")
                    return self._clean_extracted_text(text)
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    logger.info(f"Successfully extracted text from PDF using PyPDF2: {len(text)} characters")
                    return self._clean_extracted_text(text)
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {e}")
        
        # Final fallback - try to read as text (for simple PDFs)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                # Basic cleaning for PDF artifacts
                text = re.sub(r'[^\x00-\x7F]+', ' ', content)  # Remove non-ASCII
                if len(text.strip()) > 50:  # Only if we got reasonable content
                    logger.info("Extracted text using fallback text reading")
                    return self._clean_extracted_text(text)
        except Exception as e:
            logger.error(f"Fallback text extraction failed: {e}")
        
        logger.error("Failed to extract text from PDF")
        return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX extraction")
            return ""
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {len(text)} characters")
            return self._clean_extracted_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        logger.info(f"Successfully read TXT file with {encoding} encoding: {len(text)} characters")
                        return self._clean_extracted_text(text)
                except UnicodeDecodeError:
                    continue
            
            logger.error("Failed to read TXT file with any encoding")
            return ""
            
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove page numbers and headers/footers (basic patterns)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^Page \d+ of \d+$', '', text, flags=re.MULTILINE)
        
        # Clean up common resume artifacts
        text = re.sub(r'•|◦|▪|‣', '•', text)  # Normalize bullet points
        
        # Ensure text is not too long (for API limits)
        if len(text) > 5000:
            text = text[:5000] + "..."
            logger.info("Text truncated to 5000 characters for processing")
        
        return text.strip()
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about a file"""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            stat = os.stat(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            info = {
                "size": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "extension": file_ext,
                "name": Path(file_path).name,
                "is_supported": file_ext in ['.pdf', '.docx', '.txt', '.doc']
            }
            
            # Try to get page count for PDFs
            if file_ext == '.pdf' and PDF_AVAILABLE:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        info["pages"] = len(pdf_reader.pages)
                except:
                    info["pages"] = "Unknown"
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {"error": str(e)}